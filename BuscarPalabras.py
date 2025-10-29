import os
import re
import datetime
import pandas as pd
from docx import Document

# === CONFIGURACIÓN ===
#archivo_original = r"C:\Users\CIGSA\Desktop\Caso Formateado OOEE\Anexo_1 Licitación OOEE, I&C GOM MODULO 1 (29_10_25).docx"
#archivo_original = r"C:\Users\CIGSA\Desktop\Caso Formateado OOEE\Anexo_1 Licitación OOEE, I&C GOM MODULO 2 (29_10_25).docx"
archivo_original = r"C:\Users\CIGSA\Desktop\Solo Macro Word Formato\Formateados\Anexo 1 Licitación Desarrollos Verticales_RevA_23-10 1 1.docx"
#["ee","cc","contratista","codelco","det"]
palabra_clave = "adc"  # <-- Cambia esto por la palabra que quieras buscar

sensibilidad_mayusculas = False  # ← Cambia a True si quieres búsqueda estricta

# === FUNCIONES ===
def buscar_en_fragmento(fragmento, palabra, sensible):
    resultados = []
    flags = 0 if sensible else re.IGNORECASE
    for match in re.finditer(re.escape(palabra), fragmento, flags):
        inicio = max(match.start() - 10, 0)
        fin = min(match.end() + 10, len(fragmento))
        contexto = fragmento[inicio:fin]
        resultados.append({
            "Fragmento": fragmento,
            "Texto alrededor": contexto
        })
    return resultados

def procesar_texto(texto, palabra, sensible):
    resultados = []
    fragmentos = texto.split('\n')
    for frag in fragmentos:
        if re.search(re.escape(palabra), frag, 0 if sensible else re.IGNORECASE):
            resultados.extend(buscar_en_fragmento(frag, palabra, sensible))
    return resultados

# === CARGAR DOCUMENTO ===
doc = Document(archivo_original)
todos_resultados = []

# Párrafos normales
for parrafo in doc.paragraphs:
    texto = parrafo.text.strip()
    todos_resultados.extend(procesar_texto(texto, palabra_clave, sensibilidad_mayusculas))

# Texto en tablas
for tabla in doc.tables:
    for fila in tabla.rows:
        for celda in fila.cells:
            texto = celda.text.strip()
            todos_resultados.extend(procesar_texto(texto, palabra_clave, sensibilidad_mayusculas))

# === GUARDAR EN EXCEL ===
fecha_hora = datetime.datetime.now().strftime("%Y%m%d_%H%M")
nombre_excel = f"resultado_busqueda_{fecha_hora}.xlsx"
directorio = os.path.dirname(archivo_original)
ruta_excel = os.path.join(directorio, nombre_excel)

df = pd.DataFrame(todos_resultados)
df.to_excel(ruta_excel, index=False)

print(f"✅ Resultados guardados en: {ruta_excel}")
