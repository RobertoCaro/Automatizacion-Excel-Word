import os
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

# Ruta del archivo original
archivo_original = r"C:\Users\CIGSA\Desktop\Solo Macro Word Formato\Anexo 1 Licitación Desarrollos Verticales_RevA_23-10 1.docx"

# Carpeta de salida
carpeta_salida = r"C:\Users\CIGSA\Desktop\Solo Macro Word Formato\Formateados"
Path(carpeta_salida).mkdir(parents=True, exist_ok=True)

# Cargar documento
documento = Document(archivo_original)

# Parámetros de formato
ancho1 = Pt(3.5 * 28.35)       # 2 cm
ancho2 = Pt(14.5 * 28.35)      # 10 cm
ancho_total = Pt((3.5+14.5) * 28.35) # 12 cm
altura_exacta = Pt(0.4 * 28.35)  # 0.4 cm

# Contadores
total_tablas = len(documento.tables)
tablas_con_actividad = 0
tablas_formateadas = 0

# Función para desactivar autofit
def desactivar_autofit(tabla):
    tbl = tabla._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        tbl.insert(0, tbl_pr)
    tbl_width = OxmlElement('w:tblW')
    tbl_width.set(qn('w:w'), '0')
    tbl_width.set(qn('w:type'), 'auto')
    tbl_pr.append(tbl_width)

# Función para aplicar alineación vertical sin cambiar fuente
def alinear_vertical(celda):
    celda.vertical_alignment = 1  # Centrado vertical

# Procesar tablas
for tabla in documento.tables:
    encontrado = False

    # Buscar "a)" y "Actividad"
    for fila in tabla.rows:
        for celda in fila.cells:
            texto = celda.text.replace("\n", "")
            if "a)" in texto and "Actividad" in texto:
                encontrado = True
                break
        if encontrado:
            break

    if encontrado:
        tablas_con_actividad += 1
        desactivar_autofit(tabla)

        # Ajustar anchos y aplicar alineación
        for fila in tabla.rows:
            num_celdas = len(fila.cells)
            if num_celdas == 2:
                fila.cells[0].width = ancho1
                fila.cells[1].width = ancho2
            elif num_celdas == 1:
                fila.cells[0].width = ancho_total

            for celda in fila.cells:
                alinear_vertical(celda)

        # Ajustar alturas
        for fila in tabla.rows:
            texto_fila = fila.cells[0].text.replace("\n", "")
            if any(pregunta in texto_fila for pregunta in ["a)", "b)", "c)", "d)", "e)", "f)", "g)", "i)", "j)"]):
                fila.height = altura_exacta
                fila.height_rule = 2  # altura exacta

        tablas_formateadas += 1

# Guardar nuevo archivo
nombre_archivo = Path(archivo_original).stem + "_Formateado.docx"
ruta_guardado = os.path.join(carpeta_salida, nombre_archivo)
documento.save(ruta_guardado)

# Mostrar resumen
print("✅ Proceso finalizado.")
print(f"📄 Tablas totales en documento: {total_tablas}")
print(f"🔍 Tablas que contienen 'a) Actividad': {tablas_con_actividad}")
print(f"🎯 Tablas formateadas con éxito: {tablas_formateadas}")
print(f"💾 Archivo guardado como: {ruta_guardado}")