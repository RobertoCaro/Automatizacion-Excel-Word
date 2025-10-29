import os
import shutil
from zipfile import ZipFile
from lxml import etree
from docx import Document
from docx2pdf import convert
import fitz  # PyMuPDF
from openpyxl import Workbook
from collections import Counter
from docx.shared import Pt, RGBColor

# Configuración
archivo_original = r"C:\Users\CIGSA\Desktop\Solo Macro Word Formato\Formateados\Anexo 1 Licitación Desarrollos Verticales_RevA_23-10 1 1.docx"
color_objetivo = "8DB3E2"
archivo_temp_docx = archivo_original.replace(".docx", "_temp.docx")
archivo_pdf = archivo_temp_docx.replace(".docx", ".pdf")
archivo_excel = archivo_original.replace(".docx", "_reporte.xlsx")

# Paso 1: Crear copia temporal del Word
shutil.copyfile(archivo_original, archivo_temp_docx)

# Paso 2: Extraer contenido del .docx temporal
with ZipFile(archivo_temp_docx, 'r') as zip_in:
    zip_in.extractall("temp_docx")

# Paso 3: Cargar el XML principal
ruta_xml = os.path.join("temp_docx", "word", "document.xml")
parser = etree.XMLParser(remove_blank_text=True)
tree = etree.parse(ruta_xml, parser)
root = tree.getroot()

# Paso 4: Detectar tablas con celdas coloreadas
doc = Document(archivo_temp_docx)
tablas_info = []
for tabla_idx, tabla in enumerate(doc.tables):
    celdas_coloreadas = []
    fragmentos = []
    for fila in tabla.rows:
        for celda in fila.cells:
            texto = celda.text.strip()
            xml_celda = celda._tc
            shading = xml_celda.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd")
            if shading:
                fill = shading[0].get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill")
                if fill and fill.lower() != "auto" and fill.upper() != "FFFFFF":
                    celdas_coloreadas.append({
                        "texto": texto,
                        "color": fill.upper(),
                        "ok": fill.upper() == color_objetivo
                    })
                    if texto and texto not in fragmentos and len(fragmentos) < 10:
                        fragmentos.append(texto)
    if celdas_coloreadas:
        tablas_info.append({
            "tabla_idx": tabla_idx + 1,
            "fragmentos": fragmentos,
            "celdas": celdas_coloreadas,
            "marcador": f"TABLA_{tabla_idx + 1}_MARCA"
        })

# Paso 5: Insertar marcador visible en el Word temporal
doc_temp = Document(archivo_temp_docx)
for idx, tabla in enumerate(doc_temp.tables):
    marcador = f"TABLA_{idx + 1}_MARCA"
    tabla_element = tabla._element
    parent = tabla_element.getparent()
    index = parent.index(tabla_element)

    # Crear párrafo con texto blanco y tamaño mínimo
    parrafo = doc_temp.add_paragraph()
    run = parrafo.add_run(marcador)
    run.font.size = Pt(1)
    run.font.color.rgb = RGBColor(255, 255, 255)

    # Insertar el párrafo antes de la tabla
    nuevo_parrafo = run._element.getparent()
    parent.insert(index, nuevo_parrafo)

doc_temp.save(archivo_temp_docx)

# Paso 6: Convertir a PDF
convert(archivo_temp_docx, archivo_pdf)
doc_pdf = fitz.open(archivo_pdf)

# Paso 7: Buscar marcadores en el PDF
for tabla in tablas_info:
    marcador = tabla["marcador"]
    for page_num in range(len(doc_pdf)):
        page_text = doc_pdf[page_num].get_text()
        if marcador in page_text:
            tabla["pagina"] = page_num + 1
            break
    else:
        tabla["pagina"] = "?"

# Paso 8: Crear archivo Excel
wb = Workbook()
ws = wb.active
ws.title = "Reporte de Colores"
ws.append(["Tabla Nº", "Texto Celda", "Página", "Color Detectado", "¿Color OK?"])
for tabla in tablas_info:
    for celda in tabla["celdas"]:
        ws.append([
            tabla["tabla_idx"],
            celda["texto"],
            tabla["pagina"],
            celda["color"],
            "✅" if celda["ok"] else "❌"
        ])
wb.save(archivo_excel)

# Paso 9: Limpiar temporales
shutil.rmtree("temp_docx")
doc_pdf.close()
os.remove(archivo_pdf)
os.remove(archivo_temp_docx)

print(f"📊 Reporte generado: {archivo_excel}")
print(f"🔍 Tablas con celdas coloreadas detectadas: {len(tablas_info)}")